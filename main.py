from canvasapi import Canvas
import getpass
import datetime
import tqdm
import pandas as pd
import os
from annotations import get_annotations, get_urls
from utils import CanvasSession

import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
import scipy.stats as stats
from docx import Document
from docx.shared import Inches

import warnings

warnings.filterwarnings('ignore', 'SettingWithCopyWarning')
warnings.simplefilter(action='ignore', category=FutureWarning)

def main():
    print(
"""
                      ░▓▓░   ░░                             
                       ░░    ▓▓                             
                   ░░       ░██░       ░░                   
                 ░░█▒░  ░██████████░  ░▒█░░                 
            ░▒░  ░░█▒░░     ░██░      ░▒█░░ ░░▒░            
             ▒     ░░        ▓▓        ░░    ░▒             
                    ░░░      ░░      ░░░                    
                 ▓███████▓░      ░▓███████▓                 
        ░▒▒░   ▒██▒░░░  ▒██▒░  ░▒██▒░░▒ ░▒██▒   ░▒▒░        
        ░███▓ ░██░░█▒    ░██░░░░██░░▓▓░▒█░░██░ ▓███░        
        ░▓░▓█▒▒█▒░▓░      ▒██████▒░█░░█░░  ▒█▒▒█▓░▓░        
         ░ ▒█▓░██░        ▓█▒░░▒█▓░░▓▓░   ░██░▓█▓ ░         
          ░▓██░▓█▓░     ░▒██░  ░▓█▓░     ░▓█▓░██▓░          
          ░████░░███▓▓▓███▒░    ░▒███▓▓▓███░░████░          
          ▒▒▒██░  ░▒▓▓▓▒░░        ░░▒▓▓▓▒░  ░██▒▒▒          
           ░███░       ░░░░      ░░░░       ░███░           
           ▒███░     ▒██████▒  ▒██████▒     ░███▒           
          ░██▒░    ▒████████████████████▒    ░▒██░          
         ░▓█████▓███████████▒░░▒███████████▓▓████▓░         
         ░███████████████▓░      ░▒███████████████░         
         ▓██▓███░░░░░░░     ░▓▓░     ░░░░░░░███▓██▓░        
        ░▓█░▓█████▓▓▓▓▓ ░  ▒████▒  ░ ▓▓▓▓▓█████▓░█▓░        
         ▒▓░███████████░▓░░██████░░█░███████████░▓▒         
          ░ █████████████▓▒██████▒▓█████████████ ░          
            ░██▒████████████████████████████▒██░            
             ▓█░████████████████████████████░█▓░            
             ░▓░▒██████████████████████████▒░▓░             
                 ▓██▒▓████████████████▓▒██▓                 
                 ░██▒░███▓▓███████▓███▒▒██░                 
                   ▒█░▒██▓▒██████▓▒██▒░█▒                   
                    ░░░░██░▓█████░▓█░ ░░                    
                         ▒▒░▓███░▒▒░                        
                             ▒▓                                             
""" )
    print("")
    print("www.canvaswizards.org.uk")
    print("")
    print("Welcome to the Canvas Assignment Auto Moderator!")
    print("By Robert Treharne, University of Liverpool. 2024")
    print("")

    # if config.py exists, import it
    try:
        from config import CANVAS_URL, CANVAS_TOKEN
    except ImportError:
        CANVAS_URL = input('Enter your Canvas URL: ')
        print("")
        CANVAS_TOKEN = getpass.getpass('Enter your Canvas token: ')
        print("")

    # if course_id in config.py, use it
    try:
        from config import course_id
    except ImportError:
        course_id = int(input('Enter the course ID: '))
        print("")

    # if assignment_id in config.py, use it
    try:
        from config import assignment_id
    except ImportError:
        assignment_id = int(input('Enter the assignment ID: '))
        print("")

    canvas = Canvas(CANVAS_URL, CANVAS_TOKEN)

    if input("Do you want to scrape submission annotations? (y/n): ").lower() == "y":
        annotations = True
        session = CanvasSession()
    else:
        annotations = False
        session = None

    print("")

    print("Getting submissions...")
    submissions = get_submissions(canvas, course_id, assignment_id)
    print("Getting rubric...")
    rubric = get_rubric(canvas, course_id, assignment_id)
    print("Building headers...")
    header_list = get_headers(rubric, annotations)
    print("Building report...")
    report_path = build_report(canvas, course_id, assignment_id, header_list, submissions, rubric, CANVAS_URL, annotations=annotations, session=session)
    print("Moderating report...")
    print("")
    anonymise_graders = input("Do you want to anonymise graders? (y/n): ").lower() == "y"
    print("")
    generate_summary = input("Do you want to generate a moderation summary? (y/n): ").lower() == "y"
    print("")
    moderate(report_path, anonymise_graders=anonymise_graders, generate_summary=generate_summary)

def get_submissions(canvas, course_id, assignment_id):
    course = canvas.get_course(course_id)
    assignment = course.get_assignment(assignment_id)
    submissions = [x for x in assignment.get_submissions(include=["user", "submission_comments", "rubric_assessment"])]
    return submissions

def get_rubric(canvas, course_id, assignment_id):
    course = canvas.get_course(course_id)
    assignment = course.get_assignment(assignment_id)
    return assignment.rubric

def get_headers(rubric, annotations):
    rubric_rating_headers = [f"RATING_{x['description']}" for x in rubric]
    rubric_score_headers = [f"SCORE_{x['description']}" for x in rubric]

    header_list = [
        "last_name",
        "first_name",
        "sis_user_id",
        "submitted_at",
        "seconds_late",
        "status",
        "posted_at",
        "score",
        "grader",
        "comments"]
    
    if annotations:
        header_list += ["annotations"]
    
    
    header_list += ["url"]

    header_list += rubric_rating_headers + rubric_score_headers

    return header_list

def get_rubric_rating(rubric, rubric_assessment):
    """
    Retrieves the descriptions of the ratings for each rubric item in the rubric assessment.

    Parameters:
    rubric (list): A list of rubric items.
    rubric_assessment (dict): A dictionary containing the rubric assessment data.

    Returns:
    list: A list of rating descriptions for each rubric item in the rubric assessment.
    """
    ratings_list = []
    rubric_flag = False
    for item in rubric:
        rating_id = rubric_assessment[item["id"]]["rating_id"]
        for ratings in item["ratings"]:
            if ratings["id"] == rating_id:
                if ratings["description"]:
                    ratings_list.append(ratings["description"])
                    rubric_flag = True
                else:
                    ratings_list.append("")
        if rubric_flag:
            rubric_flag = False
        else:
            ratings_list.append("")
            
    return ratings_list

def get_rubric_score(rubric, rubric_assessment):
    """
    Calculates the score for each rubric item based on the rubric assessment.

    Parameters:
    rubric (list): The rubric containing the criteria and ratings.
    rubric_assessment (dict): The rubric assessment containing the rating for each rubric item.

    Returns:
    list: A list of scores for each rubric item.
    """
    ratings_list = []
    for item in rubric:
        rating_id = rubric_assessment[item["id"]]["rating_id"]
        for ratings in item["ratings"]:
            if ratings["id"] == rating_id:
                    ratings_list.append(ratings["points"])
        
    return ratings_list

def build_submission_string(canvas, header_list, rubric, submission, CANVAS_URL, course_id, assignment_id, annotations=False, session=None):
    """
    Builds a row of data for a submission in a Canvas assignment report.

    Args:
        submission (Submission): The submission object representing a student's submission.

    Returns:
        list: A list containing the row of data for the submission, including student information,
              submission details, grading information, and rubric ratings and scores.
    """
    
    sortable_name = f'{submission.user["sortable_name"]}'
    last_name, first_name = sortable_name.split(", ")
    sis_user_id = submission.user["sis_user_id"]
    submitted_at = submission.submitted_at
    seconds_late = submission.seconds_late
    status = submission.workflow_state
    posted_at = submission.posted_at
    score = submission.score

    if annotations:
        url = f"{CANVAS_URL}/courses/{course_id}/gradebook/speed_grader?assignment_id={assignment_id}&student_id={submission.user_id}"
        ann = get_annotations(session, url)
        ann = ",".join([x["comment"] for x in ann])

    url = f"{CANVAS_URL}/courses/{course_id}/gradebook/speed_grader?assignment_id={assignment_id}&student_id={submission.user_id}"

    try:
        grader = canvas.get_user(submission.grader_id).sortable_name
    except:
        grader = ""
    
    try:
        rubric_assessment = submission.rubric_assessment
    except:
        rubric_assessment = ""
        
    comments = ", ".join([f"{x["comment"]}" for x in submission.submission_comments])

    if rubric_assessment:
        rubric_rating = get_rubric_rating(rubric, rubric_assessment)
        rubric_score = get_rubric_score(rubric, rubric_assessment)
    else:
        rubric_rating = [""]*len(rubric)
        rubric_score = [""]*len(rubric)

    values = [
        last_name,
        first_name,
        sis_user_id,
        submitted_at,
        seconds_late,
        status,
        posted_at,
        score,
        grader,
        comments]
    
    if annotations:
        values += [ann]

    values += [
        url
    ]

    values += rubric_rating + rubric_score

    row = {}

    for header, value in zip(header_list, values):
        row[header] = value
        
    return row

def build_report(canvas, course_id, assignment_id, header_list, submissions, rubric, CANVAS_URL, annotations=False, session=None):
    course = canvas.get_course(course_id)
    assignment = course.get_assignment(assignment_id)
    dirname = course.course_code
    subdirname = assignment.name[:20].replace(" ", "_")

    # check if course directory exists, if not, create it
    if not os.path.exists(dirname):
        os.makedirs(dirname)

    # check if assignment directory exists, if not, create it
    if not os.path.exists(os.path.join(dirname, assignment.name[:20].replace(" ", "_"))):
        os.makedirs(os.path.join(dirname, subdirname))

    fpath = os.path.join(dirname, subdirname, f"{assignment.name[:20].replace(" ", "_")}_moderation_report.xlsx")
    
    if os.path.exists(fpath):
        data = pd.read_excel(fpath)
        for submission in tqdm.tqdm(submissions, desc=f"Writing submissions to Excel file {fpath}"):
            if submission.user["sis_user_id"] in data["sis_user_id"].values:
                continue
            else:
                submission_string = build_submission_string(canvas, header_list, rubric, submission, CANVAS_URL, course_id, assignment_id, annotations=annotations, session=session)
                new_row = pd.DataFrame(submission_string, index=[0])
                data = pd.concat([data, new_row], ignore_index=True)
                data.to_excel(fpath, index=False)
    else:
        data = pd.DataFrame()
        for submission in tqdm.tqdm(submissions, desc="Writing submissions to Excel file"):
            submission_string = build_submission_string(canvas, header_list, rubric, submission, CANVAS_URL, course_id, assignment_id, annotations=annotations, session=session)
            new_row = pd.DataFrame(submission_string, index=[0])
            data = pd.concat([data, new_row], ignore_index=True)
            data.to_excel(fpath, index=False)

    return fpath

def grader_analysis(df, fpath, label='score'):
    # Convert 'score' column to numeric type
    df[label] = pd.to_numeric(df[label], errors='coerce')

    df = df[df[label]>0]
       
    # Grader analysis
    global_median = df[label].median()
    global_mean = df[label].mean() 

    # Perform analysis for each grader
    results = []
    for grader in df['grader'].unique():
        try:
            print(grader)
            grader_scores = df[df['grader'] == grader][label]
            median_score = grader_scores.median()
            mean_score = grader_scores.mean()
            other_scores = df[df['grader'] != grader][label]
            print("grader_scores", grader_scores)
            print("other_scores", other_scores)
            U_stat, p_value = stats.mannwhitneyu(grader_scores, other_scores, alternative='two-sided')
            results.append({'Grader': grader, f'Median {label}': median_score, f'Mean {label}': mean_score, 'P-Value': p_value})
        except:
            continue


    results_df = pd.DataFrame(results)
    results_df.sort_values(f'Median {label}', inplace=True)

    # Create the horizontal boxplot
    plt.figure(figsize=(10, 15))
    box_plot = sns.boxplot(x=label, y='grader', data=df, order=results_df['Grader'], orient='h')

    # Update axis labels
    plt.xlabel(label)
    plt.ylabel('Grader')

    # Overlay the strip plot to show individual scores
    sns.stripplot(x=label, y='grader', data=df, order=results_df['Grader'], color='red', size=5, jitter=True, orient='h', alpha=0.5)

    # Add a vertical line for the global median score
    plt.axvline(x=global_median, color='green')

    # Annotate significant differences
    for i, grader in enumerate(results_df['Grader']):
        if results_df[results_df['Grader'] == grader]['P-Value'].values[0] < 0.05:
            plt.text(x=df[label].max() + 1, y=i+0.3, s=f"*", fontsize=20, color='red', verticalalignment='center')



    plt.tight_layout()
    

    # save the plot
    box_plot.figure.savefig(fpath.replace("moderation_report.xlsx", f"{label}_boxplot.png"))

    # Get graders with significant differences
    significant_graders = results_df[results_df['P-Value'] < 0.05]['Grader'].values

    return results_df, significant_graders

def count_total_words(df):
    # If annotations column exists, count the total number of words in the annotations
    if 'annotations' in df.columns:
        df['total_annotations_words'] = df['annotations'].str.split().str.len()
    else:
        df['total_annotations_words'] = 0
    

    if 'comments' in df.columns:
        df['total_comments_words'] = df['comments'].str.split().str.len()
    else:
        df['total_comments_words'] = 0

    df['total_words'] = df['total_annotations_words'] + df['total_comments_words']

    return df

def moderate(fpath, anonymise_graders=False, generate_summary=False):
    df = pd.read_excel(fpath)

    # Add column "moderate_reason" to df. Default is empty string
    df['moderation_issue'] = ''
    df['rubric_score_diff'] = 0

    # Only keep "graded" and nonzero scores
    df = df[df['status'] == 'graded']
    df = df[df['score'] > 0]

    global_median = df['score'].median()
    global_mean = df['score'].mean()

    df = count_total_words(df)

    if anonymise_graders:

        # Randomise df
        df = df.sample(frac=1)

        grader_hash = {grader: i for i, grader in enumerate(df['grader'].unique())}

        # convert grader_hash to dataframe and save
        grader_hash_df = pd.DataFrame(grader_hash.items(), columns=['grader', 'hash'])
        grader_hash_df.to_csv(fpath.replace("moderation_report.xlsx", "grader_hash.csv"), index=False)

        # Replace grader names with hash
        df['grader'] = df['grader'].map(grader_hash)

    results_df, significant_graders = grader_analysis(df, fpath, label='score')

    # For each grader, if their median score is significantly different to the global median, set moderate to True and moderate_reason to "Grader median score is significantly different to global median"
    for grader in significant_graders:
        median_score = results_df[results_df['Grader'] == grader]['Median score'].values[0]
        if median_score > global_median:
            df.loc[df['grader'] == grader, 'moderation_issue'] += f"Grader median score is significantly higher than global median, "
        else:
            df.loc[df['grader'] == grader, 'moderation_issue'] += f"Grader median score is significantly lower than global median, "

    if 'total_words' in df.columns:
        words_df, significant_graders_words = grader_analysis(df, fpath, label='total_words')

    print(significant_graders_words)


    # for all df, total values for each row in columns containing word "SCORE". 
    # If total is more than value in "score" column, set moderate to True and moderate_reason to "Final score is different to rubric total"

    df.loc[df.filter(like='SCORE').sum(axis=1) != df['score'], 'moderation_issue'] += "Final score is different to rubric total, "

    df.loc[df.filter(like='SCORE').sum(axis=1) != df['score'], 'rubric_score_diff'] = df.filter(like='SCORE').sum(axis=1) - df['score']

    df.to_excel(fpath, index=False)
    print(f"Moderated report saved as {fpath}")

    if generate_summary:
        doc = Document()

        # Add title
        doc.add_heading(f'Moderation Summary: {fpath} ', level=0)

        # Add credit
        doc.add_paragraph("This moderation summary was generated using the Canvas Assignment Auto Moderator by R. Treharne. For more information, contact R.Treharne@liverpool.ac.uk")

        # Add the boxplot
        doc.add_picture(fpath.replace("moderation_report.xlsx", "boxplot.png"), width=Inches(5))

        # Add figure caption
        doc.add_paragraph('Figure 1: Boxplot of scores by grader. Red dots indicate individual scores. Green line indicates the global median score. * indicates significant differences between grader median scores and the global median score.')

        # get graders with P-Value < 0.05
        significant_graders = results_df[results_df["P-Value"] < 0.05]

        # Format median and mean values to 2 decimal places
        significant_graders["Median Score"] = significant_graders["Median Score"].map("{:.2f}".format)
        significant_graders["Mean Score"] = significant_graders["Mean Score"].map("{:.2f}".format)

        # Format P-Value to scientific notation
        significant_graders["P-Value"] = significant_graders["P-Value"].map(lambda x: f"{x:.2e}")

        # Start a new page
        doc.add_page_break()

        # Add table title
        doc.add_paragraph('Table 1: Graders with significant differences in median scores compared to the global median score.')

        # Add significatn_graders dataframe as table to document
        t = doc.add_table(significant_graders.shape[0]+1, significant_graders.shape[1])

        # add the header rows.
        for j in range(significant_graders.shape[-1]):
            t.cell(0,j).text = significant_graders.columns[j]

        # add the rest of the data frame
        for i in range(significant_graders.shape[0]):
            for j in range(significant_graders.shape[-1]):
                t.cell(i+1,j).text = str(significant_graders.values[i,j])

        # Save the document
        doc.save(fpath.replace("moderation_report.xlsx", "moderation_summary.docx"))


if __name__ == "__main__":
    main()